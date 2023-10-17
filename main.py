import pypandoc
import os
import docx
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
import datetime
import re
from pdf2docx import Converter
from tqdm import tqdm
import yaml

today=str(datetime.date.today()).replace("-","")
excl = []
fail = []
Var = locals()
with open('./company.yaml','r',encoding='utf-8') as f:
    Company_Name = yaml.load(f.read(),Loader=yaml.FullLoader) #在此添加公司名称
Company_List = []
for i in range(len(Company_Name)+1):
    Var['doc' + str(i)] = docx.Document()
    Company_List.append(Var['doc' + str(i)])
doc_else = docx.Document()
k = doc1.styles["Heading 1"]
out_format = "docx" #文件转换输出的后缀
path = "./TS/" #把所有格式的纪要放入TS文件夹
file_list = list(os.listdir(path))
file_list.reverse()
print(file_list)

#file_name是包含后缀名的文件名
for file_name in tqdm(file_list):
    name = os.path.splitext(file_name)[0] #name是不包含后缀名的文件名
    try: #时间和来源都有
        if re.search("\d{6,8}",name):
            time = re.search("\d{6,8}",name).group()
        elif re.search("\d{2,4}Q[1,4]",name):
            time = re.search("\d{2,4}Q[1,4]",name).group()
        elif re.search("\d{1,2}\s[A-Za-z]{3,3}\s\d{4,4}",name):
            time = re.search("\d{1,2}\s[A-Za-z]{3,3}\s\d{4,4}",name).group()
        elif re.search("\d{4,4}年\d+月\d+日",name):
            time = re.search("\d{4,4}年\d+月\d+日",name).group()
        else:
            time = re.search("\d{2,4}Q[1,4]",name).group() #进入except
        source = re.match("【.*】",name).group()
        name = name.replace(time,"")
        name = name.replace(source,"")
        name = time + name + source
    except:
        try: #只有时间没有来源
            if re.search("\d{6,8}", name):
                time = re.search("\d{6,8}", name).group()
            elif re.search("\d{2,4}Q[1,4]", name):
                time = re.search("\d{2,4}Q[1,4]", name).group()
            elif re.search("\d{1,2}\s[A-Za-z]{3,3}\s\d{4,4}", name):
                time = re.search("\d{1,2}\s[A-Za-z]{3,3}\s\d{4,4}", name).group()
            elif re.search("\d{4,4}年\d+月\d+日", name):
                time = re.search("\d{4,4}年\d+月\d+日", name).group()
            else:
                time = re.search("\d{2,4}Q[1,4]", name).group()  # 进入except
            name = name.replace(time,"")
            name = time + name
        except:
            try: #只有来源没有时间
                source = re.match("【.*】", name).group()
                name = name.replace(source, "")
                name = name + source
            except Exception as e:
                excl.append(name)
                pass
    in_name = path + file_name #in_name是文件的绝对路径
    #print(file_name,".md" in file_name or ".html" in file_name)
    if ".md" in file_name or ".html" in file_name: #如果是md文件，直接读取内容，添加文件名后写入word
        if ".html" in file_name:
            pypandoc.convert_file(in_name, 'md', outputfile=path+name+".md") #如果是html文件，转换为md文件
            os.remove(in_name)
            in_name = path+name+".md"
            #print(in_name)
        with open(in_name,"r+",encoding="utf-8") as f:
            #print(".md" or ".html" in file_name, file_name)
            content = f.read()
        for j in range(len(Company_Name)):
            if Company_Name[j] in name:
                Company_List[j].add_heading(name,1)
                Company_List[j].add_paragraph(content)
                Company_List[j].add_page_break()
                f.close()
                break
            elif j == len(Company_Name)-1:
                doc_else.add_heading(name,1)
                doc_else.add_paragraph(content)
                doc_else.add_page_break()
                f.close()
                continue

    elif "docx" in file_name or ".pdf" in file_name:
        if ".pdf" in file_name:
            cv = Converter(in_name)
            cv.convert(path+name+"."+out_format)
            cv.close()
            os.remove(in_name)
            in_name = path + name + "." + out_format
        document = docx.Document(in_name)
        content = document.paragraphs
        for j in range(len(Company_Name)):
            if Company_Name[j] in name:
                Company_List[j].add_heading(name, 1)
                for para in content:
                    Company_List[j].add_paragraph(para.text)
                Company_List[j].add_page_break()
                break
            elif j == len(Company_Name) - 1:
                doc_else.add_heading(name, 1)
                for para in content:
                    doc_else.add_paragraph(para.text)
                doc_else.add_page_break()
                continue

        '''
            content = "#" + name + "#" + "\n" + content
            f.seek(0)
            f.truncate()
            f.write(content)
            f.close()
            '''
        '''
        word = pypandoc.convert_file(in_name,out_format,format=None,encoding="utf-8",outputfile=out_name)
        doc = docx.Document(out_name)
        '''
        '''
        style = doc.styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
        '''
        ''' 
        doc.add_paragraph(name,style=k)
        doc.save(out_name)
        '''
    else:
        fail.append(file_name)
        pass
for a in range(len(Company_Name)):
    Company_List[a].styles['Normal'].font.name = u"微软雅黑"
    Company_List[a].styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'),u"微软雅黑")
    Company_List[a].save(Company_Name[a]+"纪要"+today+".docx")
doc_else.save("行业纪要"+today+".docx")
for s in excl:
    print('文件名不符合规则---'+s)
for s in fail:
    print('后缀名不符合规则---'+s)